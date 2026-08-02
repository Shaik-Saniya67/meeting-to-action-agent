"""
AI Meeting-to-Action Agent
--------------------------------------
Built for: AI Builders Challenge with IBM Bob (Wildcard: Future of Work)
Author: Shaik Saniya

Takes a raw meeting transcript and turns it into a full action system:
  - Structured minutes (topics, decisions, open questions)
  - Extracted action items with owner / deadline / priority + reasoning
  - AI Clarification Agent: flags missing info instead of silently guessing
  - Risk Detector: flags risky situations (unfinalized owners, tight deadlines)
  - Conflict Detection: duplicate ownership, contradicting decisions
  - Meeting Insights Dashboard: quick stats
  - Personalized per-person follow-up emails
  - Export to PDF / DOCX / CSV / ICS calendar events
  - Action item status tracker (Pending / In Progress / Done)

AI: Google Gemini (gemini-flash-latest), JSON structured output
Storage: SQLite (local, zero-config)
Dev tool: Built with the assistance of IBM Bob
"""

import streamlit as st
import google.generativeai as genai
import sqlite3
import json
import os
import io
import csv
from datetime import datetime, timedelta

# --------------------------------------------------------------------------
# CONFIG
# --------------------------------------------------------------------------
st.set_page_config(page_title="Meeting-to-Action AI Agent", page_icon="🗒️", layout="wide")

DB_PATH = os.path.join(os.getcwd(), "meetings.db")
MODEL_NAME = "gemini-flash-latest"
USAGE_FILE = os.path.join(os.getcwd(), "host_usage.json")
HOST_QUOTA_LIMIT = 15  # AI generations per day allowed on the shared host demo key

WEEKDAYS = ["monday", "tuesday", "wednesday", "thursday", "friday", "saturday", "sunday"]

# --------------------------------------------------------------------------
# DATABASE
# --------------------------------------------------------------------------
def init_db():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("""
        CREATE TABLE IF NOT EXISTS meetings (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT,
            meeting_date TEXT,
            meeting_priority TEXT DEFAULT 'Medium',
            created_at TEXT,
            transcript TEXT,
            result_json TEXT
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS action_items (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            meeting_id INTEGER,
            owner TEXT,
            task TEXT,
            deadline TEXT,
            deadline_date TEXT,
            priority TEXT,
            priority_reason TEXT,
            missing_info TEXT,
            status TEXT DEFAULT 'Pending',
            FOREIGN KEY(meeting_id) REFERENCES meetings(id)
        )
    """)
    conn.commit()
    # lightweight migration for anyone running an older DB file
    for col, coltype in [("meeting_date", "TEXT"), ("meeting_priority", "TEXT DEFAULT 'Medium'"),
                          ("deadline_date", "TEXT"),
                          ("priority_reason", "TEXT"), ("missing_info", "TEXT"),
                          ("status", "TEXT DEFAULT 'Pending'")]:
        try:
            table = "meetings" if col in ("meeting_date", "meeting_priority") else "action_items"
            c.execute(f"ALTER TABLE {table} ADD COLUMN {col} {coltype}")
            conn.commit()
        except sqlite3.OperationalError:
            pass  # column already exists
    conn.close()


def save_meeting(title, meeting_date, meeting_priority, transcript, result):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute(
        "INSERT INTO meetings (title, meeting_date, meeting_priority, created_at, transcript, result_json) VALUES (?, ?, ?, ?, ?, ?)",
        (title, meeting_date.isoformat(), meeting_priority, datetime.now().isoformat(), transcript, json.dumps(result)),
    )
    meeting_id = c.lastrowid
    for item in result.get("action_items", []):
        deadline_date = resolve_deadline_date(item.get("deadline", ""), meeting_date)
        c.execute(
            """INSERT INTO action_items
               (meeting_id, owner, task, deadline, deadline_date, priority, priority_reason, missing_info, status)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, 'Pending')""",
            (meeting_id, item.get("owner", "Unassigned"), item.get("task", ""),
             item.get("deadline", "Not specified"),
             deadline_date.isoformat() if deadline_date else None,
             item.get("priority", "medium"), item.get("priority_reason", ""),
             json.dumps(item.get("missing_info", []))),
        )
    conn.commit()
    conn.close()
    return meeting_id


def get_all_meetings():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT id, title, created_at, meeting_priority FROM meetings ORDER BY id DESC")
    rows = c.fetchall()
    conn.close()
    return rows


def get_meeting(meeting_id):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT title, created_at, transcript, result_json, meeting_priority FROM meetings WHERE id=?", (meeting_id,))
    row = c.fetchone()
    conn.close()
    return row


def delete_meetings(meeting_ids):
    if not meeting_ids:
        return
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    placeholders = ",".join("?" for _ in meeting_ids)
    c.execute(f"DELETE FROM action_items WHERE meeting_id IN ({placeholders})", meeting_ids)
    c.execute(f"DELETE FROM meetings WHERE id IN ({placeholders})", meeting_ids)
    conn.commit()
    conn.close()


def get_all_meeting_contexts():
    """Fetch a lightweight summary of every past meeting, for the AI chat feature."""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT title, meeting_date, result_json FROM meetings ORDER BY meeting_date ASC")
    rows = c.fetchall()
    conn.close()
    return rows


def get_analytics():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT COUNT(*) FROM meetings")
    total_meetings = c.fetchone()[0]
    c.execute("SELECT COUNT(*) FROM action_items WHERE status='Done'")
    completed = c.fetchone()[0]
    c.execute("SELECT COUNT(*) FROM action_items WHERE status!='Done'")
    pending = c.fetchone()[0]
    c.execute("SELECT COUNT(*) FROM action_items WHERE priority='high' AND status!='Done'")
    urgent = c.fetchone()[0]
    c.execute("""
        SELECT owner, COUNT(*) as cnt FROM action_items
        WHERE owner IS NOT NULL AND owner != '' AND LOWER(owner) != 'unassigned'
        GROUP BY owner ORDER BY cnt DESC LIMIT 1
    """)
    top = c.fetchone()

    c.execute("SELECT result_json FROM meetings")
    all_results = [json.loads(row[0]) for row in c.fetchall()]
    conn.close()

    total_conflicts = sum(len(r.get("conflicts", [])) for r in all_results)
    total_risks = sum(len(r.get("risks", [])) for r in all_results)

    scores = [compute_effectiveness(r)[0] for r in all_results]
    meeting_health = round(sum(scores) / len(scores)) if scores else 0

    all_items = [i for r in all_results for i in r.get("action_items", [])]
    total_items_all = len(all_items)
    if total_items_all:
        assigned = sum(1 for i in all_items if not _is_unassigned(i))
        dated = sum(1 for i in all_items if not _has_no_deadline(i))
        productivity = round(((assigned + dated) / (2 * total_items_all)) * 100)
    else:
        productivity = 0

    tracked = completed + pending
    completion_rate = round((completed / tracked) * 100) if tracked else 0

    return {
        "total_meetings": total_meetings,
        "completed": completed,
        "pending": pending,
        "urgent": urgent,
        "most_active": top[0] if top else "N/A",
        "most_active_count": top[1] if top else 0,
        "conflicts": total_conflicts,
        "risks": total_risks,
        "meeting_health": meeting_health,
        "productivity": productivity,
        "completion_rate": completion_rate,
    }


def get_open_action_items(only_unresolved=False):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    query = """
        SELECT action_items.id, meetings.title, action_items.owner, action_items.task,
               action_items.deadline, action_items.priority, action_items.status,
               action_items.deadline_date
        FROM action_items JOIN meetings ON action_items.meeting_id = meetings.id
    """
    if only_unresolved:
        query += " WHERE action_items.status != 'Done'"
    query += " ORDER BY (action_items.status='Done') ASC, action_items.priority DESC"
    c.execute(query)
    rows = c.fetchall()
    conn.close()
    return rows


def update_status(item_id, status):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("UPDATE action_items SET status=? WHERE id=?", (status, item_id))
    conn.commit()
    conn.close()


# --------------------------------------------------------------------------
# HOST DEMO KEY — DAILY QUOTA GUARD
# --------------------------------------------------------------------------
def _read_host_usage():
    today = datetime.now().date().isoformat()
    data = {"date": today, "count": 0}
    if os.path.exists(USAGE_FILE):
        try:
            with open(USAGE_FILE, "r") as f:
                stored = json.load(f)
            if stored.get("date") == today:
                data = stored
        except Exception:
            pass
    return data


def host_quota_remaining():
    data = _read_host_usage()
    return max(0, HOST_QUOTA_LIMIT - data.get("count", 0))


def _consume_host_quota():
    data = _read_host_usage()
    data["count"] = data.get("count", 0) + 1
    try:
        with open(USAGE_FILE, "w") as f:
            json.dump(data, f)
    except Exception:
        pass
    return data["count"]


def guard_host_quota(is_host_key):
    """Call right before any Gemini call. Returns True if OK to proceed.
    Only meters usage when the shared host key is the one being used —
    visitors using their own key are never limited."""
    if not is_host_key:
        return True
    if host_quota_remaining() <= 0:
        st.error("⚠️ The host's shared demo quota is used up for today. "
                 "Please add your own free Gemini API key in the sidebar to continue "
                 "(get one at aistudio.google.com/apikey).")
        return False
    _consume_host_quota()
    return True


# --------------------------------------------------------------------------
# MEETING EFFECTIVENESS SCORE
# --------------------------------------------------------------------------
def _is_unassigned(item):
    return item.get("owner", "").strip().lower() in ("", "unassigned")


def _has_no_deadline(item):
    return item.get("deadline", "Not specified").strip().lower() in ("", "not specified", "tbd", "n/a")


def compute_effectiveness(result):
    """Start at 100 and subtract for every open question, conflict, risk,
    unassigned task, and task missing a deadline. Floors at 0."""
    items = result.get("action_items", [])
    breakdown = {
        "open_questions": len(result.get("open_questions", [])),
        "conflicts": len(result.get("conflicts", [])),
        "risks": len(result.get("risks", [])),
        "unassigned": sum(1 for i in items if _is_unassigned(i)),
        "no_deadline": sum(1 for i in items if _has_no_deadline(i)),
    }
    score = 100
    score -= breakdown["open_questions"] * 10
    score -= breakdown["conflicts"] * 10
    score -= breakdown["risks"] * 5
    score -= breakdown["unassigned"] * 5
    score -= breakdown["no_deadline"] * 3
    score = max(0, score)
    return score, breakdown


def effectiveness_explanation(score, breakdown, result):
    """Plain-language explanation built from the same evidence as the score
    (no extra API call needed — reuses the Gemini-extracted fields)."""
    fixes = []
    if breakdown["unassigned"]:
        task_name = "the open task"
        for i in result.get("action_items", []):
            if _is_unassigned(i):
                task_name = f"'{i.get('task', 'the open task')}'"
                break
        fixes.append(f"assigning an owner for {task_name}")
    if breakdown["conflicts"]:
        fixes.append("resolving the flagged conflict")
    if breakdown["risks"]:
        fixes.append("addressing the flagged risk")
    if breakdown["open_questions"]:
        fixes.append("closing out the open questions")
    if breakdown["no_deadline"]:
        fixes.append("setting deadlines on the tasks missing one")

    if score >= 90:
        opener = "This meeting was well organized and clearly executed"
    elif score >= 70:
        opener = "This meeting is well organized"
    elif score >= 40:
        opener = "This meeting made good progress"
    else:
        opener = "This meeting surfaced good discussion"

    if fixes:
        return f"{opener}, but {' and '.join(fixes[:2])} will improve execution."
    return f"{opener} — decisions, owners, and deadlines are all in place."


def build_recommendations(result):
    """Assemble AI Project Manager recommendations from the fields Gemini
    already extracted in one pass: conflicts (with their AI resolution),
    risks, missing owners, missing deadlines, and open questions."""
    recs = []
    for c in result.get("conflicts", []):
        if c.get("recommendation"):
            recs.append(c["recommendation"])
    for r in result.get("risks", []):
        if r.get("description"):
            recs.append(f"Address risk: {r['description']}")
    for i in result.get("action_items", []):
        if _is_unassigned(i):
            recs.append(f"Assign an owner to: {i.get('task', 'this task')}.")
        if _has_no_deadline(i):
            recs.append(f"Clarify the deadline for: {i.get('task', 'this task')}.")
    for q in result.get("open_questions", []):
        recs.append(f"Resolve open question: {q}")

    seen, unique = set(), []
    for r in recs:
        if r not in seen:
            seen.add(r)
            unique.append(r)
    return unique[:8]


# --------------------------------------------------------------------------
# DEADLINE RESOLUTION (weekday names -> actual dates, relative to meeting date)
# --------------------------------------------------------------------------
def resolve_deadline_date(deadline_text, meeting_date):
    if not deadline_text:
        return None
    text = deadline_text.strip().lower()
    if text in ("not specified", "tbd", "n/a", ""):
        return None
    if text == "today":
        return meeting_date
    if text == "tomorrow":
        return meeting_date + timedelta(days=1)
    for i, day in enumerate(WEEKDAYS):
        if day in text:
            days_ahead = (i - meeting_date.weekday() + 7) % 7
            days_ahead = 7 if days_ahead == 0 else days_ahead  # assume next occurrence, not today
            return meeting_date + timedelta(days=days_ahead)
    # try common absolute date formats
    for fmt in ("%Y-%m-%d", "%d-%m-%Y", "%d/%m/%Y", "%B %d, %Y", "%b %d, %Y", "%B %d", "%b %d"):
        try:
            parsed = datetime.strptime(deadline_text.strip(), fmt)
            year = parsed.year if parsed.year != 1900 else meeting_date.year
            return datetime(year, parsed.month, parsed.day).date()
        except ValueError:
            continue
    return None


# --------------------------------------------------------------------------
# GEMINI AGENT PROMPTS
# --------------------------------------------------------------------------
EXTRACTION_PROMPT = """You are a precise business meeting analyst and clarification agent. Read the meeting transcript below and extract structured information.

Return ONLY valid JSON matching this exact schema, no markdown fences, no commentary:

{{
  "meeting_summary": "2-3 sentence high level summary of what the meeting was about",
  "topics_discussed": ["topic 1", "topic 2"],
  "decisions_made": ["decision 1", "decision 2"],
  "open_questions": ["unresolved question 1"],
  "action_items": [
    {{
      "owner": "person name or 'Unassigned' if genuinely unclear",
      "task": "what they need to do",
      "deadline": "date, weekday name, or 'Not specified'",
      "priority": "high" or "medium" or "low",
      "priority_reason": "one short sentence explaining why this priority was chosen",
      "missing_info": ["specific clarifying question, e.g. 'Who owns this task?' or 'What is the expected deadline?'"]
    }}
  ],
  "conflicts": [
    {{"type": "duplicate_ownership" or "contradiction", "description": "plain language description of the conflict and why it matters",
      "recommendation": "a specific, actionable resolution suggestion, reasoned from evidence in the transcript (e.g. who already started the work, what was committed to earlier, what minimizes rework before the deadline)"}}
  ],
  "risks": [
    {{"description": "plain language description of a risky situation, e.g. unfinalized ownership on a near-term deadline, or an unresolved decision blocking other work"}}
  ]
}}

Rules:
- Infer the owner of an action item even if only implied (e.g. "I'll take care of it" -> assign to the speaker who said it, if identifiable). Only use "Unassigned" if genuinely no one is implicated.
- missing_info should list a clarifying question ONLY for fields that are genuinely ambiguous or absent (no owner, no deadline, vague task description). If everything is clear, return an empty list for that item.
- priority is "high" if there's explicit urgency (near deadline, "asap", "critical", blocks other work), "medium" for normal work with a deadline, "low" for open-ended or minor items. Always include a one-sentence priority_reason.
- Flag duplicate_ownership when two people appear assigned to the same or overlapping task.
- Flag contradiction when a later statement in the transcript conflicts with an earlier decision.
- For every conflict, always include a concrete "recommendation" — don't just describe the problem, propose the specific resolution a project lead would reasonably make, grounded in transcript evidence (prior commitments, who already started work, what minimizes rework before the deadline).
- Flag risks for situations like: an urgent task with no owner, a tight deadline with unresolved details, or a decision that blocks other action items but remains open.
- If there are no conflicts or risks, return empty lists.

Transcript:
\"\"\"
{transcript}
\"\"\"
"""

EMAIL_PROMPT = """Based on this meeting analysis JSON, write a concise, professional follow-up email recapping the meeting for the whole team.
Include: a short intro line, key decisions, and a clear action items list with owners and deadlines.
Keep it under 200 words. Return plain text only, no JSON, no markdown fences.

Meeting analysis:
{analysis}
"""

PERSONALIZED_EMAIL_PROMPT = """Based on this meeting analysis JSON, write short, personalized follow-up emails — one for each unique action item owner (skip "Unassigned").
Each email should remind that person specifically of what they're responsible for and by when, in a friendly professional tone, under 100 words each.

Return ONLY valid JSON, no markdown fences, matching this schema:
{{
  "emails": [
    {{"owner": "name", "subject": "short subject line", "body": "email body text"}}
  ]
}}

Meeting analysis:
{analysis}
"""


def get_model(api_key):
    genai.configure(api_key=api_key)
    return genai.GenerativeModel(MODEL_NAME)


def _clean_json_text(text):
    text = text.strip()
    if text.startswith("```"):
        text = text.split("```")[1]
        if text.startswith("json"):
            text = text[4:]
    return text.strip()


def extract_minutes(api_key, transcript):
    model = get_model(api_key)
    prompt = EXTRACTION_PROMPT.format(transcript=transcript)
    response = model.generate_content(
        prompt,
        generation_config={"response_mime_type": "application/json", "temperature": 0.3},
        request_options={"timeout": 90},
    )
    return json.loads(_clean_json_text(response.text))


def draft_email(api_key, analysis):
    model = get_model(api_key)
    prompt = EMAIL_PROMPT.format(analysis=json.dumps(analysis))
    response = model.generate_content(
        prompt,
        generation_config={"temperature": 0.5},
        request_options={"timeout": 90},
    )
    return response.text.strip()


def draft_personalized_emails(api_key, analysis):
    model = get_model(api_key)
    prompt = PERSONALIZED_EMAIL_PROMPT.format(analysis=json.dumps(analysis))
    response = model.generate_content(
        prompt,
        generation_config={"response_mime_type": "application/json", "temperature": 0.5},
        request_options={"timeout": 40},
    )
    return json.loads(_clean_json_text(response.text)).get("emails", [])


CHAT_PROMPT = """You are an AI co-worker who has attended every meeting on this team and remembers all of them.
Answer the user's question using ONLY the meeting history provided below. Be specific — name people, tasks, and dates when relevant.
If the answer isn't in the meeting history, say so plainly rather than guessing.

Meeting history (chronological):
{history}

Question: {question}

Answer concisely, in plain text (no markdown fences).
"""


def ask_about_meetings(api_key, question, meeting_contexts):
    history_parts = []
    for title, mdate, result_json in meeting_contexts:
        r = json.loads(result_json)
        items_str = "; ".join(
            f"{i.get('owner')}: {i.get('task')} (deadline {i.get('deadline')}, priority {i.get('priority')})"
            for i in r.get("action_items", [])
        )
        history_parts.append(
            f"Meeting: {title} ({mdate})\n"
            f"Summary: {r.get('meeting_summary', '')}\n"
            f"Decisions: {'; '.join(r.get('decisions_made', []))}\n"
            f"Action items: {items_str}\n"
        )
    history = "\n---\n".join(history_parts) if history_parts else "No meetings recorded yet."

    model = get_model(api_key)
    prompt = CHAT_PROMPT.format(history=history, question=question)
    response = model.generate_content(
        prompt,
        generation_config={"temperature": 0.3},
        request_options={"timeout": 40},
    )
    return response.text.strip()


def transcribe_audio(api_key, audio_bytes, mime_type):
    """Transcribe a meeting recording using Gemini's native audio understanding
    (no separate speech-to-text service needed)."""
    model = get_model(api_key)
    audio_part = {"mime_type": mime_type, "data": audio_bytes}
    prompt = (
        "Transcribe this meeting recording as accurately as possible. "
        "Label speakers if you can distinguish different voices (e.g. 'Speaker 1:', 'Speaker 2:'), "
        "or use inferred names if someone addresses another person by name. "
        "Return plain text only — no commentary, no markdown."
    )
    response = model.generate_content(
        [audio_part, prompt],
        request_options={"timeout": 120},
    )
    return response.text.strip()


# --------------------------------------------------------------------------
# EXPORT HELPERS
# --------------------------------------------------------------------------
def export_csv(result):
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["Owner", "Task", "Deadline", "Priority", "Priority Reason"])
    for item in result.get("action_items", []):
        writer.writerow([item.get("owner"), item.get("task"), item.get("deadline"),
                          item.get("priority"), item.get("priority_reason", "")])
    return buf.getvalue().encode("utf-8")


def export_pdf(result, title):
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(0, 10, title, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 11)
    pdf.ln(2)

    def section(heading, lines):
        pdf.set_font("Helvetica", "B", 13)
        pdf.multi_cell(0, 8, heading, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", "", 11)
        for line in lines:
            pdf.multi_cell(0, 7, f"- {line}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(2)

    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 7, result.get("meeting_summary", ""), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)
    section("Decisions Made", result.get("decisions_made", []))
    section("Open Questions", result.get("open_questions", []))

    pdf.set_font("Helvetica", "B", 13)
    pdf.multi_cell(0, 8, "Action Items", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 11)
    for item in result.get("action_items", []):
        pdf.multi_cell(0, 7, f"- [{item.get('priority','').upper()}] {item.get('owner')}: "
                              f"{item.get('task')} (by {item.get('deadline')})",
                        new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)

    conflicts = result.get("conflicts", [])
    if conflicts:
        lines = []
        for c in conflicts:
            lines.append(c.get("description", ""))
            if c.get("recommendation"):
                lines.append(f"Recommendation: {c.get('recommendation')}")
        section("Conflicts Detected", lines)
    risks = result.get("risks", [])
    if risks:
        section("Risks Detected", [r.get("description", "") for r in risks])

    return bytes(pdf.output())


def export_docx(result, title):
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(result.get("meeting_summary", ""))

    doc.add_heading("Decisions Made", level=2)
    for d in result.get("decisions_made", []):
        doc.add_paragraph(d, style="List Bullet")

    doc.add_heading("Open Questions", level=2)
    for q in result.get("open_questions", []):
        doc.add_paragraph(q, style="List Bullet")

    doc.add_heading("Action Items", level=2)
    for item in result.get("action_items", []):
        doc.add_paragraph(
            f"[{item.get('priority','').upper()}] {item.get('owner')}: {item.get('task')} "
            f"(by {item.get('deadline')})", style="List Bullet"
        )

    conflicts = result.get("conflicts", [])
    if conflicts:
        doc.add_heading("Conflicts Detected", level=2)
        for c in conflicts:
            doc.add_paragraph(c.get("description", ""), style="List Bullet")
            if c.get("recommendation"):
                doc.add_paragraph(f"Recommendation: {c.get('recommendation')}", style="List Bullet 2")

    risks = result.get("risks", [])
    if risks:
        doc.add_heading("Risks Detected", level=2)
        for r in risks:
            doc.add_paragraph(r.get("description", ""), style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_ics(result, meeting_date):
    lines = ["BEGIN:VCALENDAR", "VERSION:2.0", "PRODID:-//Meeting-to-Action Agent//EN"]
    for item in result.get("action_items", []):
        ddate = resolve_deadline_date(item.get("deadline", ""), meeting_date)
        if not ddate:
            continue
        dt = ddate.strftime("%Y%m%d")
        uid = f"{dt}-{abs(hash(item.get('task','')))}@meeting-agent"
        lines += [
            "BEGIN:VEVENT",
            f"UID:{uid}",
            f"DTSTAMP:{datetime.now().strftime('%Y%m%dT%H%M%SZ')}",
            f"DTSTART;VALUE=DATE:{dt}",
            f"SUMMARY:{item.get('task','Task')} ({item.get('owner','Unassigned')})",
            f"DESCRIPTION:Priority: {item.get('priority','')}. {item.get('priority_reason','')}",
            "END:VEVENT",
        ]
    lines.append("END:VCALENDAR")
    return "\r\n".join(lines).encode("utf-8")


# --------------------------------------------------------------------------
# UI — DESIGN SYSTEM (CSS + reusable card renderers)
# --------------------------------------------------------------------------
init_db()

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&family=Lexend:wght@500;600;700;800&display=swap');

:root{
  --ink:#1a1d29; --muted:#6b7280; --line:rgba(17,24,39,0.08);
  --violet:#7c5cff; --blue:#3b82f6; --grad:linear-gradient(135deg,#6d5bff 0%,#3b82f6 100%);
  --grad-soft:linear-gradient(135deg,rgba(124,92,255,0.10) 0%,rgba(59,130,246,0.10) 100%);
  --green:#12b76a; --amber:#f79009; --red:#f04438; --card-bg:rgba(255,255,255,0.72);
}
html, body, [class*="css"], .stMarkdown, p, div, span { font-family:'Inter',sans-serif; }
h1,h2,h3,h4,.hero-title{ font-family:'Lexend',sans-serif; }

.stApp{
  background:
    radial-gradient(1200px 500px at 10% -10%, rgba(124,92,255,0.10), transparent 60%),
    radial-gradient(1000px 500px at 100% 0%, rgba(59,130,246,0.10), transparent 55%),
    #f7f8fc;
}
section[data-testid="stSidebar"]{
  background:linear-gradient(180deg,#171227 0%, #1c1836 100%);
  border-right:1px solid rgba(255,255,255,0.06);
}
section[data-testid="stSidebar"] *{ color:#eceaf6 !important; }
section[data-testid="stSidebar"] .stTextInput input{
  background:rgba(255,255,255,0.07); border:1px solid rgba(255,255,255,0.12);
  border-radius:12px; color:#fff !important;
}
section[data-testid="stSidebar"] .stRadio > div{ gap:6px; }
section[data-testid="stSidebar"] .stRadio label{
  padding:9px 14px; border-radius:12px; transition:all .15s ease; width:100%;
}
section[data-testid="stSidebar"] .stRadio label:hover{ background:rgba(255,255,255,0.08); }
section[data-testid="stSidebar"] hr{ border-color:rgba(255,255,255,0.10); }

/* Hero */
.hero{
  background:var(--grad); border-radius:20px; padding:34px 38px; margin-bottom:24px;
  box-shadow:0 12px 32px -12px rgba(109,91,255,0.45); position:relative; overflow:hidden;
}
.hero::after{
  content:""; position:absolute; right:-60px; top:-60px; width:220px; height:220px;
  background:rgba(255,255,255,0.10); border-radius:50%;
}
.hero-title{ color:#fff; font-size:30px; font-weight:800; margin:0 0 6px 0; letter-spacing:-0.3px; }
.hero-sub{ color:rgba(255,255,255,0.88); font-size:14.5px; margin:0; max-width:720px; line-height:1.5; }
.hero-pill{
  display:inline-block; background:rgba(255,255,255,0.16); color:#fff; font-size:12px;
  font-weight:600; padding:4px 12px; border-radius:999px; margin-bottom:12px; letter-spacing:.2px;
}

/* Section headers */
.sec-head{ display:flex; align-items:center; gap:10px; margin:26px 0 14px 0; }
.sec-head .ic{
  width:34px; height:34px; border-radius:10px; background:var(--grad-soft);
  display:flex; align-items:center; justify-content:center; font-size:16px;
}
.sec-head h3{ margin:0; font-size:18px; font-weight:700; color:var(--ink); }

/* Glass metric cards */
.metric-grid{ display:grid; grid-template-columns:repeat(6,1fr); gap:14px; margin-bottom:6px; }
@media(max-width:1100px){ .metric-grid{ grid-template-columns:repeat(3,1fr);} }
.metric-card{
  background:var(--card-bg); backdrop-filter:blur(10px); border:1px solid var(--line);
  border-radius:18px; padding:18px 16px; box-shadow:0 4px 20px -8px rgba(17,24,39,0.08);
  transition:transform .18s ease, box-shadow .18s ease;
}
.metric-card:hover{ transform:translateY(-3px); box-shadow:0 14px 28px -12px rgba(109,91,255,0.28); }
.metric-card .m-ic{ font-size:20px; margin-bottom:8px; display:block; }
.metric-card .m-val{ font-size:26px; font-weight:800; color:var(--ink); line-height:1.1; }
.metric-card .m-label{ font-size:12.5px; color:var(--muted); margin-top:3px; font-weight:500; }

/* Score ring */
.score-wrap{ display:flex; align-items:center; gap:26px; }
.ring{
  width:128px; height:128px; border-radius:50%; display:flex; align-items:center; justify-content:center;
  background:conic-gradient(var(--ring-color,#7c5cff) calc(var(--pct,0)*1%), #edeafc calc(var(--pct,0)*1%));
  flex:none;
}
.ring-inner{
  width:100px; height:100px; border-radius:50%; background:#fff; display:flex; flex-direction:column;
  align-items:center; justify-content:center; box-shadow:inset 0 0 0 1px var(--line);
}
.ring-inner .num{ font-size:26px; font-weight:800; color:var(--ink); line-height:1; }
.ring-inner .lbl{ font-size:10.5px; color:var(--muted); margin-top:2px; }
.check-list .row{ font-size:14px; margin-bottom:6px; color:var(--ink); }
.check-list .ok{ color:var(--green); font-weight:600; }
.check-list .warn{ color:var(--amber); font-weight:600; }

/* Badges */
.badge{ display:inline-block; font-size:11.5px; font-weight:700; padding:3px 10px; border-radius:999px; letter-spacing:.2px; }
.badge-high{ background:rgba(240,68,56,0.12); color:#c0281c; }
.badge-medium{ background:rgba(247,144,9,0.14); color:#a15c00; }
.badge-low{ background:rgba(18,183,106,0.13); color:#0a7a48; }
.badge-pending{ background:rgba(107,114,128,0.14); color:#4b5563; }
.badge-inprogress{ background:rgba(59,130,246,0.14); color:#1d4ed8; }
.badge-done{ background:rgba(18,183,106,0.14); color:#0a7a48; }
.chip{ display:inline-block; font-size:11.5px; font-weight:600; padding:3px 10px; border-radius:999px;
  background:#f1f0fb; color:#5a4fcf; margin-left:6px; }

/* Action item / conflict / risk / recommendation cards */
.a-card{
  background:#fff; border:1px solid var(--line); border-radius:16px; padding:16px 18px; margin-bottom:12px;
  box-shadow:0 2px 10px -6px rgba(17,24,39,0.06); transition:box-shadow .15s ease, transform .15s ease;
}
.a-card:hover{ box-shadow:0 10px 22px -10px rgba(109,91,255,0.25); transform:translateY(-1px); }
.a-top{ display:flex; align-items:center; justify-content:space-between; gap:10px; flex-wrap:wrap; }
.a-left{ display:flex; align-items:center; gap:12px; }
.avatar{
  width:34px; height:34px; border-radius:50%; background:var(--grad); color:#fff; font-weight:700;
  font-size:13px; display:flex; align-items:center; justify-content:center; flex:none;
}
.a-task{ font-size:14.5px; color:var(--ink); font-weight:600; margin-top:8px; line-height:1.45; }
.a-reason{ font-size:12.5px; color:var(--muted); margin-top:4px; }
.a-owner{ font-weight:700; font-size:13.5px; color:var(--ink); }
.a-deadline{ font-size:12px; color:var(--muted); }

.warn-card{ border-radius:16px; padding:14px 18px; margin-bottom:10px; border:1px solid; display:flex; gap:12px; }
.warn-card .ic{ font-size:18px; }
.conflict-card{ background:rgba(247,144,9,0.06); border-color:rgba(247,144,9,0.28); }
.risk-card{ background:rgba(240,68,56,0.06); border-color:rgba(240,68,56,0.25); }
.rec-card{ background:rgba(18,183,106,0.06); border-color:rgba(18,183,106,0.25); }
.warn-title{ font-weight:700; font-size:13.5px; color:var(--ink); margin-bottom:3px; }
.warn-body{ font-size:13.5px; color:#374151; line-height:1.5; }
.warn-rec{ font-size:13px; color:#0a7a48; margin-top:6px; font-weight:500; }

/* Buttons */
.stButton>button{
  border-radius:12px !important; font-weight:600 !important; border:1px solid var(--line) !important;
  transition:all .15s ease !important;
}
.stButton>button[kind="primary"]{
  background:var(--grad) !important; border:none !important;
  box-shadow:0 8px 18px -8px rgba(109,91,255,0.55) !important;
}
.stButton>button:hover{ transform:translateY(-1px); }
.stDownloadButton>button{ border-radius:12px !important; font-weight:600 !important; }

/* Containers / expanders */
div[data-testid="stExpander"]{
  border-radius:16px !important; border:1px solid var(--line) !important; overflow:hidden;
  box-shadow:0 2px 10px -6px rgba(17,24,39,0.05);
}
div[data-testid="stVerticalBlockBorderWrapper"]{ border-radius:16px !important; }

/* Chat bubbles */
[data-testid="stChatMessage"]{
  border-radius:16px !important; padding:4px 6px !important; margin-bottom:6px !important;
}

/* Alerts */
div[data-testid="stAlert"]{ border-radius:14px !important; }
</style>
""", unsafe_allow_html=True)


def badge_html(kind, text):
    cls = {"high": "badge-high", "medium": "badge-medium", "low": "badge-low",
           "Pending": "badge-pending", "In Progress": "badge-inprogress", "Done": "badge-done"}.get(text, kind)
    return f'<span class="badge {cls}">{text.upper() if kind == "priority" else text}</span>'


def avatar_initials(name):
    name = (name or "?").strip()
    parts = [p for p in name.replace("_", " ").split() if p]
    if not parts:
        return "?"
    if len(parts) == 1:
        return parts[0][:2].upper()
    return (parts[0][0] + parts[-1][0]).upper()


def render_metric_grid(metrics):
    """metrics: list of (icon, label, value) tuples."""
    cards = "".join(
        f'<div class="metric-card"><span class="m-ic">{icon}</span>'
        f'<div class="m-val">{value}</div><div class="m-label">{label}</div></div>'
        for icon, label, value in metrics
    )
    st.markdown(f'<div class="metric-grid">{cards}</div>', unsafe_allow_html=True)


def render_score_ring(score, color="#7c5cff"):
    st.markdown(f"""
    <div class="ring" style="--pct:{score}; --ring-color:{color};">
      <div class="ring-inner"><div class="num">{score}</div><div class="lbl">/ 100</div></div>
    </div>
    """, unsafe_allow_html=True)


def section_header(icon, title):
    st.markdown(f'<div class="sec-head"><div class="ic">{icon}</div><h3>{title}</h3></div>', unsafe_allow_html=True)


def render_action_item_card(item):
    priority = item.get("priority", "medium")
    owner = item.get("owner", "Unassigned") or "Unassigned"
    st.markdown(f"""
    <div class="a-card">
      <div class="a-top">
        <div class="a-left">
          <div class="avatar">{avatar_initials(owner)}</div>
          <div><div class="a-owner">{owner}</div><div class="a-deadline">📅 {item.get('deadline', 'Not specified')}</div></div>
        </div>
        {badge_html('priority', priority)}
      </div>
      <div class="a-task">{item.get('task', '')}</div>
      <div class="a-reason">💭 {item.get('priority_reason', '')}</div>
    </div>
    """, unsafe_allow_html=True)
    missing = item.get("missing_info", [])
    if missing:
        with st.container(border=True):
            st.markdown("⚠️ **Missing Information**")
            for q in missing:
                st.markdown(f"- {q}")


def render_warn_card(css_class, icon, title, body, recommendation=None):
    rec_html = f'<div class="warn-rec">🤖 {recommendation}</div>' if recommendation else ""
    st.markdown(f"""
    <div class="warn-card {css_class}">
      <div class="ic">{icon}</div>
      <div><div class="warn-title">{title}</div><div class="warn-body">{body}</div>{rec_html}</div>
    </div>
    """, unsafe_allow_html=True)


with st.sidebar:
    st.markdown("""
    <div style="display:flex;align-items:center;gap:10px;margin:2px 0 18px 0;">
      <div style="width:38px;height:38px;border-radius:11px;background:linear-gradient(135deg,#7c5cff,#3b82f6);
                  display:flex;align-items:center;justify-content:center;font-size:19px;">🗒️</div>
      <div>
        <div style="font-weight:800;font-size:15px;color:#fff;line-height:1.2;">Meeting-to-Action</div>
        <div style="font-size:11px;color:rgba(255,255,255,0.55);">AI Agent · Gemini</div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div style="font-size:11px;font-weight:700;letter-spacing:.06em;color:rgba(255,255,255,0.5);margin-bottom:6px;">API SETUP</div>', unsafe_allow_html=True)
    user_provided_key = st.text_input(
        "Gemini API Key", type="password", label_visibility="collapsed",
        placeholder="Gemini API Key (optional)",
        help="Optional — leave blank to use the host's demo key. Get your own free at aistudio.google.com/apikey"
    )
    host_key = st.secrets.get("GEMINI_API_KEY", "")
    api_key = user_provided_key.strip() or host_key
    using_host_key = bool(not user_provided_key.strip() and host_key)

    if user_provided_key.strip():
        st.caption("✅ Using your own API key")
    elif host_key:
        remaining = host_quota_remaining()
        st.caption(f"ℹ️ Using host demo key — **{remaining}/{HOST_QUOTA_LIMIT}** generations left today.")
    else:
        st.caption("⚠️ No API key set — enter one above to use the app.")

    st.markdown('<hr>', unsafe_allow_html=True)
    st.markdown('<div style="font-size:11px;font-weight:700;letter-spacing:.06em;color:rgba(255,255,255,0.5);margin-bottom:6px;">NAVIGATE</div>', unsafe_allow_html=True)
    page = st.radio("Navigate", ["New Meeting", "Past Meetings", "Open Action Items", "AI Chat", "Analytics"],
                     label_visibility="collapsed")

st.markdown("""
<div class="hero">
  <div class="hero-pill">✨ Powered by Gemini · Built with IBM Bob</div>
  <div class="hero-title">🗒️ AI Meeting-to-Action Agent</div>
  <p class="hero-sub">Paste a meeting transcript and get structured minutes, clarification questions, risk flags,
  prioritized action items, insights, exports, and calendar events — automatically.</p>
</div>
""", unsafe_allow_html=True)

PRIORITY_BADGE = {"high": "🔴", "medium": "🟡", "low": "🟢"}
STATUS_OPTIONS = ["Pending", "In Progress", "Done"]

# --------------------------------------------------------------------------
# PAGE: NEW MEETING
# --------------------------------------------------------------------------
if page == "New Meeting":
    section_header("📝", "New Meeting")
    with st.container(border=True):
        col_a, col_b, col_c = st.columns([2, 1, 1])
        with col_a:
            title = st.text_input("Meeting title", placeholder="e.g. Weekly Sync - July 16")
        with col_b:
            meeting_date = st.date_input("Meeting date", value=datetime.now().date())
        with col_c:
            meeting_priority = st.selectbox("Meeting priority", ["Low", "Medium", "High"], index=1)

        transcript = st.text_area("Paste meeting transcript", height=260,
                                   value=st.session_state.get("transcribed_text", ""),
                                   placeholder="Smran: Let's finalize the dashboard by Friday...\nRahul: I'll handle the backend...")

    with st.expander("🎤 Or upload a meeting recording instead"):
        audio_file = st.file_uploader("Upload audio (mp3, wav, m4a)", type=["mp3", "wav", "m4a", "ogg"])
        if audio_file is not None:
            if st.button("🎙️ Transcribe audio"):
                if not api_key:
                    st.error("Please enter your Gemini API key in the sidebar first.")
                elif not guard_host_quota(using_host_key):
                    pass
                else:
                    with st.spinner("Transcribing audio... this can take a minute for longer recordings."):
                        try:
                            mime_map = {"mp3": "audio/mpeg", "wav": "audio/wav",
                                        "m4a": "audio/mp4", "ogg": "audio/ogg"}
                            ext = audio_file.name.split(".")[-1].lower()
                            mime_type = mime_map.get(ext, "audio/mpeg")
                            transcribed = transcribe_audio(api_key, audio_file.read(), mime_type)
                            st.session_state["transcribed_text"] = transcribed
                            st.success("Transcribed! Review it below, then click Generate Minutes.")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Transcription failed: {e}")

    SAMPLE_TRANSCRIPT = """Smran: Let's finalize the sales dashboard by Friday, that's the priority this week.
Rahul: I'll handle the backend API integration.
Priya: I can also take the backend integration if Rahul is busy with the deployment.
Smran: Good, let's also decide on the color scheme - I think dark navy fits our brand.
Rahul: Actually I already committed to using a light theme last week, so let's stick with that instead.
Priya: I'll prepare the demo video, need it done by Thursday since the submission is Friday.
Smran: Sounds good, let's also loop in the design team about the icons, that's still open.
"""

    sample_col, run_col = st.columns([1, 1])
    with sample_col:
        if st.button("🧪 Load sample transcript", use_container_width=True):
            st.session_state["transcribed_text"] = SAMPLE_TRANSCRIPT
            st.rerun()
    with run_col:
        run = st.button("⚡ Generate Minutes", type="primary", use_container_width=True)

    if run:
        if not api_key:
            st.error("Please enter your Gemini API key in the sidebar first.")
        elif not transcript.strip():
            st.error("Please paste a transcript or load the sample.")
        elif not guard_host_quota(using_host_key):
            pass
        else:
            with st.spinner("Analyzing transcript..."):
                try:
                    result = extract_minutes(api_key, transcript)
                    email = draft_email(api_key, result)
                    result["_email_draft"] = email
                    meeting_id = save_meeting(title or "Untitled Meeting", meeting_date, meeting_priority, transcript, result)
                    st.session_state["last_result"] = result
                    st.session_state["last_meeting_id"] = meeting_id
                    st.session_state["last_meeting_date"] = meeting_date
                    st.session_state["last_title"] = title or "Untitled Meeting"
                    st.success(f"Minutes generated and saved (Meeting #{meeting_id})")
                except Exception as e:
                    st.error(f"Something went wrong: {e}")

    result = st.session_state.get("last_result")
    if result:
        meeting_date = st.session_state.get("last_meeting_date", datetime.now().date())
        title = st.session_state.get("last_title", "Meeting")

        # ---- Insights Dashboard ----
        section_header("📊", "Meeting Insights")
        items = result.get("action_items", [])
        unassigned = sum(1 for i in items if i.get("owner", "").lower() in ("unassigned", ""))
        with_deadline = sum(1 for i in items if i.get("deadline", "Not specified").lower() != "not specified")
        render_metric_grid([
            ("🎯", "Total Tasks", len(items)),
            ("✅", "Decisions", len(result.get("decisions_made", []))),
            ("❓", "Open Questions", len(result.get("open_questions", []))),
            ("⚠️", "Conflicts", len(result.get("conflicts", []))),
            ("👤", "Unassigned", unassigned),
            ("📅", "With Deadlines", with_deadline),
        ])

        # ---- Meeting Effectiveness Score ----
        score, breakdown = compute_effectiveness(result)
        section_header("⭐", "Meeting Effectiveness Score")
        ring_color = "#12b76a" if score >= 80 else ("#f79009" if score >= 50 else "#f04438")
        sc1, sc2 = st.columns([1, 2])
        with sc1:
            render_score_ring(score, ring_color)
        with sc2:
            def cl(ok, ok_text, warn_text):
                return f'<div class="row"><span class="ok">✔ {ok_text}</span></div>' if ok else \
                       f'<div class="row"><span class="warn">⚠ {warn_text}</span></div>'
            rows = ""
            rows += cl(bool(result.get("decisions_made")),
                       f"{len(result.get('decisions_made', []))} Decisions Recorded", "No Decisions Recorded")
            rows += cl(breakdown["unassigned"] == 0, "Owners Assigned",
                       f"{breakdown['unassigned']} Unassigned Task{'s' if breakdown['unassigned'] != 1 else ''}")
            rows += cl(breakdown["no_deadline"] == 0, "Deadlines Present",
                       f"{breakdown['no_deadline']} Task{'s' if breakdown['no_deadline'] != 1 else ''} without Deadline")
            if breakdown["open_questions"]:
                rows += cl(False, "", f"{breakdown['open_questions']} Open Question{'s' if breakdown['open_questions'] != 1 else ''}")
            if breakdown["conflicts"]:
                rows += cl(False, "", f"{breakdown['conflicts']} Conflict{'s' if breakdown['conflicts'] != 1 else ''}")
            if breakdown["risks"]:
                rows += cl(False, "", f"{breakdown['risks']} Risk{'s' if breakdown['risks'] != 1 else ''}")
            st.markdown(f'<div class="check-list">{rows}</div>', unsafe_allow_html=True)
            st.caption(f"🤖 {effectiveness_explanation(score, breakdown, result)}")

        # ---- AI Recommendations ----
        recs = build_recommendations(result)
        if recs:
            section_header("💡", "AI Recommendations")
            for r in recs:
                render_warn_card("rec-card", "💡", "Suggested action", r)

        section_header("📋", "Summary")
        st.write(result.get("meeting_summary", ""))

        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**✅ Decisions Made**")
            for d in result.get("decisions_made", []):
                st.markdown(f"- {d}")
        with col2:
            st.markdown("**❓ Open Questions**")
            for q in result.get("open_questions", []):
                st.markdown(f"- {q}")

        # ---- Action Items + Clarification Agent ----
        section_header("🎯", "Action Items")
        for item in items:
            render_action_item_card(item)

        # ---- Conflicts + AI Resolution Suggestions ----
        conflicts = result.get("conflicts", [])
        if conflicts:
            section_header("⚠️", "Conflicts Detected")
            for c in conflicts:
                render_warn_card("conflict-card", "⚠️", c.get("type", "conflict"),
                                  c.get("description", ""), c.get("recommendation"))
        else:
            st.info("No conflicts detected in this meeting.")

        # ---- Risks ----
        risks = result.get("risks", [])
        if risks:
            section_header("🧠", "Risks Detected")
            for r in risks:
                render_warn_card("risk-card", "🧠", "Risk flagged", r.get("description", ""))

        # ---- Execution Plan (grouped by resolved deadline date) ----
        section_header("🗓️", "Execution Plan")
        today = meeting_date
        buckets = {}
        no_date = []
        for item in items:
            ddate = resolve_deadline_date(item.get("deadline", ""), meeting_date)
            if ddate:
                buckets.setdefault(ddate, []).append(item)
            else:
                no_date.append(item)
        if buckets:
            for ddate in sorted(buckets.keys()):
                if ddate == today:
                    label = f"Today ({ddate.strftime('%b %d')})"
                elif ddate == today + timedelta(days=1):
                    label = f"Tomorrow ({ddate.strftime('%b %d')})"
                else:
                    label = ddate.strftime("%A, %b %d")
                st.markdown(f"**{label}**")
                for item in buckets[ddate]:
                    badge = PRIORITY_BADGE.get(item.get("priority", "medium"), "🟡")
                    st.markdown(f"　{badge} {item.get('owner')} → {item.get('task')}")
        if no_date:
            st.markdown("**No deadline set**")
            for item in no_date:
                st.markdown(f"　⚪ {item.get('owner')} → {item.get('task')}")

        # ---- Workload Distribution ----
        if items:
            section_header("👥", "Workload Distribution")
            workload = {}
            for item in items:
                owner = item.get("owner", "Unassigned") or "Unassigned"
                workload[owner] = workload.get(owner, 0) + 1
            total_tasks = sum(workload.values())
            sorted_workload = sorted(workload.items(), key=lambda x: -x[1])
            max_count = sorted_workload[0][1]
            for owner, count in sorted_workload:
                bar_col, num_col = st.columns([6, 1])
                bar_col.progress(count / max_count, text=owner)
                num_col.markdown(f"**{count}**")

            top_owner, top_count = sorted_workload[0]
            if top_owner.lower() != "unassigned" and total_tasks > 0 and (top_count / total_tasks) >= 0.4:
                pct = round(top_count / total_tasks * 100)
                st.warning(f"⚠ {top_owner} has {pct}% of all tasks.")
                other_owners = [o for o, _ in sorted_workload if o != top_owner and o.lower() != "unassigned"]
                if other_owners:
                    least_owner = sorted(other_owners, key=lambda o: workload[o])[0]
                    for item in items:
                        if item.get("owner") == top_owner:
                            st.caption(f"🤖 AI Suggestion: Move '{item.get('task')}' to {least_owner}.")
                            break

        # ---- Due Soon ----
        if items:
            section_header("📅", "Due Soon")
            soon_buckets = {"Today": [], "Tomorrow": [], "This Week": [], "Next Week": [], "Later": []}
            for item in items:
                ddate = resolve_deadline_date(item.get("deadline", ""), meeting_date)
                if not ddate:
                    continue
                delta = (ddate - meeting_date).days
                if delta <= 0:
                    soon_buckets["Today"].append(item)
                elif delta == 1:
                    soon_buckets["Tomorrow"].append(item)
                elif delta <= 7:
                    soon_buckets["This Week"].append(item)
                elif delta <= 14:
                    soon_buckets["Next Week"].append(item)
                else:
                    soon_buckets["Later"].append(item)
            bucket_icons = {"Today": "🔴", "Tomorrow": "🟡", "This Week": "🟢", "Next Week": "📅", "Later": "⚪"}
            if any(soon_buckets.values()):
                for label, bucket_items in soon_buckets.items():
                    if bucket_items:
                        st.markdown(f"**{bucket_icons[label]} {label}**")
                        for item in bucket_items:
                            st.markdown(f"　{item.get('owner')} → {item.get('task')}")
            else:
                st.caption("No dated tasks yet.")

        # ---- Follow-up Email ----
        section_header("✉️", "Follow-up Email Draft")
        st.text_area("Copy this email", result.get("_email_draft", ""), height=180, label_visibility="collapsed")

        # ---- Interactive follow-up actions ----
        section_header("🔄", "What would you like to do next?")
        with st.container(border=True):
            b1, b2, b3, b4, b5 = st.columns(5)

            if b1.button("✍️ Personalized Emails", use_container_width=True):
                if guard_host_quota(using_host_key):
                    with st.spinner("Writing personalized emails..."):
                        try:
                            emails = draft_personalized_emails(api_key, result)
                            st.session_state["personalized_emails"] = emails
                        except Exception as e:
                            st.error(f"Couldn't generate personalized emails: {e}")

            pdf_bytes = None
            try:
                pdf_bytes = export_pdf(result, title)
            except Exception as e:
                b2.caption(f"PDF unavailable: {e}")
            if pdf_bytes:
                b2.download_button("📄 PDF", pdf_bytes, file_name=f"{title}.pdf", mime="application/pdf",
                                    use_container_width=True)

            docx_bytes = None
            try:
                docx_bytes = export_docx(result, title)
            except Exception as e:
                b3.caption(f"DOCX unavailable: {e}")
            if docx_bytes:
                b3.download_button("📝 DOCX", docx_bytes, file_name=f"{title}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True)

            csv_bytes = export_csv(result)
            b4.download_button("📊 CSV", csv_bytes, file_name=f"{title}_action_items.csv", mime="text/csv",
                                use_container_width=True)

            ics_bytes = export_ics(result, meeting_date)
            b5.download_button("📆 Calendar", ics_bytes, file_name=f"{title}.ics", mime="text/calendar",
                                use_container_width=True)

        if st.session_state.get("personalized_emails"):
            section_header("📧", "Personalized Emails")
            for e in st.session_state["personalized_emails"]:
                with st.expander(f"To {e.get('owner')} — {e.get('subject')}"):
                    st.write(e.get("body", ""))

# --------------------------------------------------------------------------
# PAGE: PAST MEETINGS
# --------------------------------------------------------------------------
elif page == "Past Meetings":
    meetings = get_all_meetings()
    if not meetings:
        st.info("No meetings saved yet. Generate one from 'New Meeting'.")
    else:
        section_header("🗂️", "Meeting Timeline")
        MEETING_PRIORITY_BADGE = {"High": "🔴", "Medium": "🟡", "Low": "🟢"}

        if "selected_meetings" not in st.session_state:
            st.session_state["selected_meetings"] = set()
        if "confirm_delete" not in st.session_state:
            st.session_state["confirm_delete"] = False

        search_term = st.text_input("Search meetings", placeholder="🔍 Search by meeting title...",
                                     label_visibility="collapsed")
        meetings = [m for m in meetings if search_term.strip().lower() in m[1].lower()] if search_term.strip() else meetings

        top = st.columns([3, 1, 1])
        selected = st.session_state["selected_meetings"]
        top[0].caption(f"{len(selected)} selected" if selected else f"{len(meetings)} meeting(s) — select below to delete.")
        if top[1].button("Select all"):
            st.session_state["selected_meetings"] = {m[0] for m in meetings}
            st.rerun()
        if top[2].button("Clear selection"):
            st.session_state["selected_meetings"] = set()
            st.session_state["confirm_delete"] = False
            st.rerun()

        if not meetings:
            st.info("No meetings match your search.")

        for mid, mtitle, created, mpriority in meetings:
            badge = MEETING_PRIORITY_BADGE.get(mpriority, "🟡")
            row = st.columns([0.4, 9.6])
            checked = row[0].checkbox("", value=mid in selected, key=f"sel_{mid}", label_visibility="collapsed")
            if checked and mid not in selected:
                selected.add(mid)
                st.rerun()
            elif not checked and mid in selected:
                selected.discard(mid)
                st.session_state["confirm_delete"] = False
                st.rerun()

            with row[1].expander(f"{badge} {mtitle} — {created[:16]}"):
                title, created_at, transcript, result_json, _mp = get_meeting(mid)
                result = json.loads(result_json)
                st.write(result.get("meeting_summary", ""))
                st.markdown("**Action items:**")
                for item in result.get("action_items", []):
                    render_action_item_card(item)
                if result.get("risks"):
                    st.markdown("**Risks:**")
                    for r in result["risks"]:
                        render_warn_card("risk-card", "🧠", "Risk flagged", r.get("description", ""))
                if result.get("conflicts"):
                    st.markdown("**Conflicts:**")
                    for c in result["conflicts"]:
                        render_warn_card("conflict-card", "⚠️", c.get("type", "conflict"),
                                          c.get("description", ""), c.get("recommendation"))
                if "_email_draft" in result:
                    st.text_area("Email draft", result["_email_draft"], height=150, key=f"email_{mid}")

        if selected:
            st.divider()
            if not st.session_state["confirm_delete"]:
                if st.button(f"🗑️ Delete {len(selected)} selected meeting(s)", type="primary"):
                    st.session_state["confirm_delete"] = True
                    st.rerun()
            else:
                st.warning(f"Delete {len(selected)} meeting(s) and their action items? This cannot be undone.")
                cc1, cc2 = st.columns(2)
                if cc1.button("Yes, delete permanently"):
                    delete_meetings(list(selected))
                    st.session_state["selected_meetings"] = set()
                    st.session_state["confirm_delete"] = False
                    st.success("Deleted.")
                    st.rerun()
                if cc2.button("Cancel"):
                    st.session_state["confirm_delete"] = False
                    st.rerun()

# --------------------------------------------------------------------------
# PAGE: OPEN ACTION ITEMS
# --------------------------------------------------------------------------
elif page == "Open Action Items":
    section_header("🗃️", "Open Action Items")
    only_unresolved = st.checkbox("Show only unresolved issues", value=True)
    items = get_open_action_items(only_unresolved=only_unresolved)
    if not items:
        st.info("No action items yet.")
    else:
        header = st.columns([2, 3, 1.5, 1, 1.5])
        header[0].markdown("**Owner**")
        header[1].markdown("**Task**")
        header[2].markdown("**Deadline**")
        header[3].markdown("**Priority**")
        header[4].markdown("**Status**")
        st.markdown('<hr style="margin:4px 0 10px 0;">', unsafe_allow_html=True)

        for item_id, mtitle, owner, task, deadline, priority, status, deadline_date in items:
            cols = st.columns([2, 3, 1.5, 1, 1.5])
            cols[0].markdown(
                f'<div style="display:flex;align-items:center;gap:8px;"><div class="avatar" '
                f'style="width:26px;height:26px;font-size:11px;">{avatar_initials(owner)}</div>'
                f'<b>{owner}</b></div>', unsafe_allow_html=True)
            cols[1].markdown(f"{'~~' + task + '~~' if status == 'Done' else task}")
            cols[2].markdown(f'<span class="chip">📅 {deadline}</span>', unsafe_allow_html=True)
            cols[3].markdown(badge_html("priority", priority), unsafe_allow_html=True)
            new_status = cols[4].selectbox("", STATUS_OPTIONS, index=STATUS_OPTIONS.index(status) if status in STATUS_OPTIONS else 0,
                                            key=f"status_{item_id}", label_visibility="collapsed")
            if new_status != status:
                update_status(item_id, new_status)
                st.rerun()
            st.markdown('<hr style="margin:6px 0;opacity:.5;">', unsafe_allow_html=True)

# --------------------------------------------------------------------------
# PAGE: AI CHAT WITH PREVIOUS MEETINGS
# --------------------------------------------------------------------------
elif page == "AI Chat":
    section_header("💬", "Ask about your meetings")
    st.caption("e.g. \"What tasks are still pending?\", \"Who is responsible for the dashboard?\", \"What decisions did we make last week?\"")

    contexts = get_all_meeting_contexts()
    if not contexts:
        st.info("No meetings recorded yet — generate one first so there's history to ask about.")
    else:
        if "chat_log" not in st.session_state:
            st.session_state["chat_log"] = []

        for role, msg in st.session_state["chat_log"]:
            with st.chat_message(role):
                st.write(msg)

        question = st.chat_input("Ask about your meeting history...")
        if question:
            if not api_key:
                st.error("Please enter your Gemini API key in the sidebar first.")
            elif not guard_host_quota(using_host_key):
                pass
            else:
                st.session_state["chat_log"].append(("user", question))
                with st.spinner("Checking meeting history..."):
                    try:
                        answer = ask_about_meetings(api_key, question, contexts)
                    except Exception as e:
                        answer = f"Something went wrong: {e}"
                st.session_state["chat_log"].append(("assistant", answer))
                st.rerun()

# --------------------------------------------------------------------------
# PAGE: ANALYTICS
# --------------------------------------------------------------------------
elif page == "Analytics":
    section_header("📊", "Meeting Analytics Dashboard")
    stats = get_analytics()

    render_metric_grid([
        ("🗓️", "Meetings", stats["total_meetings"]),
        ("🎯", "Open Tasks", stats["pending"]),
        ("✅", "Completed", stats["completed"]),
        ("🔴", "High Priority", stats["urgent"]),
        ("⚠️", "Conflicts", stats["conflicts"]),
        ("🧠", "Risks", stats["risks"]),
    ])

    section_header("📈", "Performance")
    hc1, hc2 = st.columns([1, 2])
    with hc1:
        health_color = "#12b76a" if stats["meeting_health"] >= 80 else ("#f79009" if stats["meeting_health"] >= 50 else "#f04438")
        render_score_ring(stats["meeting_health"], health_color)
        st.caption("Overall Meeting Health")
    with hc2:
        st.markdown(f"**Productivity — {stats['productivity']}%**")
        st.progress(stats["productivity"] / 100)
        st.markdown(f"**Completion Rate — {stats['completion_rate']}%**")
        st.progress(stats["completion_rate"] / 100)

    if stats["most_active"] != "N/A":
        st.caption(f"🏆 Most active: **{stats['most_active']}** ({stats['most_active_count']} tasks)")

    if stats["total_meetings"] == 0:
        st.info("No meetings yet — generate one from 'New Meeting' to populate analytics.")
